from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_apunte_masivo():
    doc = Document()
    
    # Configuración Global de Fuentes
    estilo_normal = doc.styles['Normal']
    estilo_normal.font.name = 'Google Sans'
    
    estilo_h1 = doc.styles['Heading 1']
    estilo_h1.font.name = 'Google Sans'
    estilo_h1.font.color.rgb = RGBColor(30, 64, 175) # Azul oscuro institucional

    # Portada del Apunte
    titulo = doc.add_heading('ROBÓTICA: Tecnologías para la Automatización', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.name = 'Google Sans'
    
    subtitulo = doc.add_paragraph('Apunte Completo - Cinemática, Transformaciones Homogéneas y D-H')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].font.size = Pt(14)
    subtitulo.runs[0].font.italic = True
    
    doc.add_page_break()

    # MEGA-DICCIONARIO: Transcripción íntegra de las 23 diapositivas
    diapositivas = [
        {
            "titulo": "1. Posición y Rotación (2D)",
            "texto": "Desplazamiento P(px, py)\n\n"
                     "Px = px + pu · cos α - pv · sen α\n"
                     "Py = py + pu · sen α + pv · cos α\n\n"
                     "Expresado matricialmente:\n"
                     "| Px |   | cos α  -sen α | | pu |   | px |\n"
                     "| Py | = | sen α   cos α | | pv | + | py |\n\n"
                     "Vector Traslación: P = [px, py]^T\n"
                     "Matriz de Rotación: R = | ix  iy |\n"
                     "                        | jx  jy |"
        },
        {
            "titulo": "2. Orientación de un Cuerpo Rígido",
            "texto": "[INSERTAR IMAGEN: Gráfico del Sistema Móvil UVW respecto al Fijo XYZ]\n\n"
                     "Componentes ortogonales:\n"
                     "ix = cos(ix, iu) = ix · iu\n"
                     "iy = cos(ix, iv) = ix · iv\n"
                     "iz = cos(ix, iw) = ix · iw\n\n"
                     "Componentes de la matriz de rotación general:\n"
                     "| ix·iu  ix·iv  ix·iw |\n"
                     "| jx·iu  jx·iv  jx·iw |\n"
                     "| kx·iu  kx·iv  kx·iw |"
        },
        {
            "titulo": "3. Posición y Rotación 3D",
            "texto": "Vector Traslación: P = [px, py, pz]^T\n\n"
                     "Matriz Rotación 3x3:\n"
                     "R = | nx  ox  ax |\n"
                     "    | ny  oy  ay |\n"
                     "    | nz  oz  az |\n\n"
                     "Donde los vectores unitarios corresponden a:\n"
                     "n = normal\n"
                     "o = orientación\n"
                     "a = aproximación"
        },
        {
            "titulo": "4. Rotación sobre el Eje X de α°",
            "texto": "[INSERTAR IMAGEN: Gráfico de rotación sobre X con regla de la mano derecha]\n\n"
                     "Matriz Rx,α:\n"
                     "| 1      0         0    |\n"
                     "| 0    cos α    -sen α  |\n"
                     "| 0    sen α     cos α  |"
        },
        {
            "titulo": "5. Rotación sobre el Eje Y de β°",
            "texto": "[INSERTAR IMAGEN: Gráfico de rotación sobre Y]\n\n"
                     "Matriz Ry,β:\n"
                     "|  cos β     0     sen β |\n"
                     "|    0       1       0   |\n"
                     "| -sen β     0     cos β |"
        },
        {
            "titulo": "6. Rotación sobre el Eje Z de θ°",
            "texto": "[INSERTAR IMAGEN: Gráfico de rotación sobre Z]\n\n"
                     "Matriz Rz,θ:\n"
                     "| cos θ   -sen θ     0 |\n"
                     "| sen θ    cos θ     0 |\n"
                     "|   0        0       1 |"
        },
        {
            "titulo": "7. Composición de Rotaciones Sucesivas",
            "texto": "El orden en el que se realizan las rotaciones SÍ altera el resultado.\n"
                     "La multiplicación de matrices no es conmutativa.\n\n"
                     "Rx,α · Ry,β · Rz,θ  ≠  Rz,θ · Ry,β · Rx,α"
        },
        {
            "titulo": "8. Ejemplo Práctico 7",
            "texto": "Rotar 90° respecto al eje Z al punto P1 = [2, 0, 0]^T\n\n"
                     "P2 = Rz(90) · P1\n\n"
                     "| cos 90  -sen 90   0 |   | 2 |\n"
                     "| sen 90   cos 90   0 | · | 0 |\n"
                     "|   0        0      1 |   | 0 |\n\n"
                     "Resolviendo los senos y cosenos:\n"
                     "|  0  -1   0 |   | 2 |   | 0 |\n"
                     "|  1   0   0 | · | 0 | = | 2 |\n"
                     "|  0   0   1 |   | 0 |   | 0 |\n\n"
                     "P2 = [0, 2, 0]^T\n\n"
                     "[INSERTAR IMAGEN: Gráfico 3D del Ejemplo 7]"
        },
        {
            "titulo": "9. Ejemplo Práctico 8",
            "texto": "Rotar 90° respecto al eje X y luego 90° respecto al eje Y al punto P1 = [2, 0, 0]^T\n\n"
                     "P2 = Ry(90) · (Rx(90) · P1)\n\n"
                     "Paso 1: Rx(90) · P1\n"
                     "| 1     0        0    |   | 2 |   | 2 |\n"
                     "| 0   cos 90  -sen 90 | · | 0 | = | 0 |\n"
                     "| 0   sen 90   cos 90 |   | 0 |   | 0 |\n\n"
                     "Paso 2: Ry(90) · Resultado\n"
                     "|  cos 90   0   sen 90 |   | 2 |   |  0 |\n"
                     "|    0      1     0    | · | 0 | = |  0 |\n"
                     "| -sen 90   0   cos 90 |   | 0 |   | -2 |\n\n"
                     "P2 = [0, 0, -2]^T\n\n"
                     "[INSERTAR IMAGEN: Gráfico 3D del Ejemplo 8]"
        },
        {
            "titulo": "10. Matriz de Transformación Homogénea (Forma General)",
            "texto": "Permite determinar la posición y orientación de un objeto con respecto a un sistema de ejes coordenados dados. Describe la traslación y rotación en un solo conjunto de valores (Matriz 4x4).\n\n"
                     "T = |  R_3x3         P_3x1      |\n"
                     "    |  Persp_1x3     Escala_1x1 |\n\n"
                     "• Matriz R_3x3: Orientación del sistema.\n"
                     "• Vector P_3x1: Descentramiento entre orígenes.\n"
                     "• Perspectiva: Típicamente [0, 0, 0].\n"
                     "• Escala: Típicamente 1."
        },
        {
            "titulo": "11. Significado Geométrico de la Matriz Homogénea",
            "texto": "Si se tiene un sistema de coordenadas móvil O'UVW rotado y trasladado respecto a un fijo OXYZ, la matriz T define el sistema O'UVW.\n\n"
                     "Representación en el sistema fijo:\n"
                     "P_xyz = T · P_uvw\n\n"
                     "[INSERTAR IMAGEN: Ejes OXYW y O'UVW]"
        },
        {
            "titulo": "12. Ejemplo 13 - Transformación Homogénea",
            "texto": "Trasladar al punto P(1, 2, 3) un vector u = (4, 5, 6)^T.\n\n"
                     "Vector en formato homogéneo: u = [4, 5, 6, 1]^T\n\n"
                     "Matriz de traslación T(1,2,3):\n"
                     "| 1  0  0  1 |\n"
                     "| 0  1  0  2 |\n"
                     "| 0  0  1  3 |\n"
                     "| 0  0  0  1 |\n\n"
                     "v = T · u\n"
                     "v = | 1  0  0  1 |   | 4 |   | 5 |\n"
                     "    | 0  1  0  2 | · | 5 | = | 7 |\n"
                     "    | 0  0  1  3 |   | 6 |   | 9 |\n"
                     "    | 0  0  0  1 |   | 1 |   | 1 |\n\n"
                     "v = [5, 7, 9, 1]^T"
        },
        {
            "titulo": "13. Ejemplo 14 - Rotación Homogénea",
            "texto": "Rotar un vector u = (4, 5, 6)^T un ángulo de 90° respecto al eje Z.\n\n"
                     "Matriz de rotación homogénea RotZ(90):\n"
                     "| cos 90  -sen 90  0  0 |\n"
                     "| sen 90   cos 90  0  0 |\n"
                     "|   0        0     1  0 |\n"
                     "|   0        0     0  1 |\n\n"
                     "Sustituyendo valores:\n"
                     "|  0 -1  0  0 |   | 4 |   | -5 |\n"
                     "|  1  0  0  0 | · | 5 | = |  4 |\n"
                     "|  0  0  1  0 |   | 6 |   |  6 |\n"
                     "|  0  0  0  1 |   | 1 |   |  1 |\n\n"
                     "v = [-5, 4, 6, 1]^T"
        },
        {
            "titulo": "14. Cinemática del Robot",
            "texto": "Cálculo de la posición y orientación de un robot según sus coordenadas articulares.\n\n"
                     "Método de Denavit y Hartenberg (1955):\n"
                     "Método matricial sistemático para describir los elementos de un robot respecto a un sistema de referencia fijo usando 4 parámetros (θ, d, a, α)."
        },
        {
            "titulo": "15. Pasos del Algoritmo de Denavit-Hartenberg",
            "texto": "1. Identificación de eslabones, articulaciones y ejes articulares.\n"
                     "2. Asociar un sistema de coordenadas (S) a cada eslabón i.\n"
                     "3. Asignar el origen del sistema.\n"
                     "4. Asignar el sentido de Xi.\n"
                     "5. Obtener los parámetros de D-H.\n"
                     "6. Obtener matrices individuales Ai.\n"
                     "7. Multiplicar para obtener la matriz global H."
        },
        {
            "titulo": "16. Los 4 Parámetros D-H",
            "texto": "• ai (Longitud del eslabón): Distancia desde Zi-1 a Zi medida a lo largo de Xi.\n"
                     "• αi (Torsión del eslabón): Ángulo entre Zi-1 a Zi medido alrededor de Xi.\n"
                     "• di (Desplazamiento articular): Distancia desde Xi-1 a Xi medida a lo largo de Zi-1.\n"
                     "• θi (Ángulo articular): Ángulo entre Xi-1 y Xi medido alrededor de Zi-1.\n\n"
                     "[INSERTAR IMAGEN: Eslabones y articulaciones D-H]"
        },
        {
            "titulo": "17. Reglas para la Asignación de Ejes (Casos 1, 2 y 3)",
            "texto": "Caso 1: Zi-1 y Zi no son coplanares.\n"
                     "La distancia más corta se define por ai. Xi es perpendicular a ambos.\n"
                     "[INSERTAR IMAGEN: Gráfico Caso 1]\n\n"
                     "Caso 2: Los ejes Zi-1 y Zi se cortan.\n"
                     "Xi se elige perpendicular al plano formado por ambos. ai = 0.\n"
                     "[INSERTAR IMAGEN: Gráfico Caso 2]\n\n"
                     "Caso 3: El eje Zi-1 es paralelo al eje Zi.\n"
                     "Xi se elige a lo largo de la perpendicular común. La distancia es ai.\n"
                     "[INSERTAR IMAGEN: Gráfico Caso 3]"
        },
        {
            "titulo": "18. Obtención de la Matriz Transformación Ai",
            "texto": "Ai = Rot(Z, θi) · Tras(Z, di) · Tras(X, ai) · Rot(X, αi)\n\n"
                     "Desarrollo matricial expandido de Ai:\n"
                     "| cos θi  -cos αi·sen θi   sen αi·sen θi   ai·cos θi |\n"
                     "| sen θi   cos αi·cos θi  -sen αi·cos θi   ai·sen θi |\n"
                     "|   0          sen αi          cos αi          di    |\n"
                     "|   0             0               0             1    |"
        },
        {
            "titulo": "19. Multiplicación de la Cadena Cinemática",
            "texto": "Paso 7: Multiplicar las matrices Ai para obtener la matriz resultante Hn0 (Transformación del efector final respecto a la base).\n\n"
                     "Hn0 = A1 · A2 · A3 · ... · An"
        },
        {
            "titulo": "20. Ejemplo 25 - Producto de Matrices D-H",
            "texto": "Dadas A1 y A2, hallar H = A1 · A2\n\n"
                     "A1 = \n"
                     "|  c1  -s1   0   a1·c1 |\n"
                     "|  s1   c1   0   a1·s1 |\n"
                     "|   0    0   1     0   |\n"
                     "|   0    0   0     1   |\n\n"
                     "A2 = \n"
                     "|  c2  -s2   0   a2·c2 |\n"
                     "|  s2   c2   0   a2·s2 |\n"
                     "|   0    0   1     0   |\n"
                     "|   0    0   0     1   |\n\n"
                     "H = A1 · A2 = \n"
                     "| c12  -s12   0   a1·c1 + a2·c12 |\n"
                     "| s12   c12   0   a1·s1 + a2·s12 |\n"
                     "|  0     0    1          0       |\n"
                     "|  0     0    0          1       |"
        },
        {
            "titulo": "21. Resolución: Robot Planar (2 GDL)",
            "texto": "[INSERTAR IMAGEN: Esquema físico del Robot Planar]\n\n"
                     "Tabla D-H:\n"
                     "Artic. |  θi |  di |  ai |  αi \n"
                     "   1   |  θ1 |  0  |  a1 |  0  \n"
                     "   2   |  θ2 |  0  |  a2 |  0  \n\n"
                     "Las matrices A1 y A2 coinciden con el Ejemplo 25.\n"
                     "Ecuaciones Cinemáticas Finales del Efector:\n"
                     "x = a1·cos(θ1) + a2·cos(θ1+θ2)\n"
                     "y = a1·sen(θ1) + a2·sen(θ1+θ2)"
        },
        {
            "titulo": "22. Resolución: Robot Cilíndrico (3 GDL) - Asignación",
            "texto": "[INSERTAR IMAGEN: Esquema físico del Robot Cilíndrico con Ejes Z0, Z1, Z2]\n\n"
                     "Tabla de Parámetros D-H Obtenida:\n"
                     "Artic. |  θi |  di |  ai |  αi \n"
                     "   1   |  θ1 |  d1 |  0  |  0  \n"
                     "   2   |  0  |  d2 |  0  | -90 \n"
                     "   3   |  0  |  d3 |  0  |  0  "
        },
        {
            "titulo": "23. Resolución: Robot Cilíndrico - Desarrollo Matricial",
            "texto": "Matriz A1 (Rotación en base):\n"
                     "A1 = Rot(z,θ1) · Tras(z,d1) · Tras(x,0) · Rot(x,0)\n"
                     "| c1  -s1   0   0 |\n"
                     "| s1   c1   0   0 |\n"
                     "|  0    0   1  d1 |\n"
                     "|  0    0   0   1 |\n\n"
                     "Matriz A2 (Traslación vertical d2):\n"
                     "A2 = Rot(z,0) · Tras(z,d2) · Tras(x,0) · Rot(x,-90)\n"
                     "|  1   0   0   0 |\n"
                     "|  0   0   1   0 |\n"
                     "|  0  -1   0  d2 |\n"
                     "|  0   0   0   1 |\n\n"
                     "Matriz A3 (Traslación radial d3):\n"
                     "A3 = Rot(z,0) · Tras(z,d3) · Tras(x,0) · Rot(x,0)\n"
                     "|  1   0   0   0 |\n"
                     "|  0   1   0   0 |\n"
                     "|  0   0   1  d3 |\n"
                     "|  0   0   0   1 |\n\n"
                     "Matriz Final H = A1 · A2 · A3\n"
                     "| c1   0   -s1   -s1·d3 |\n"
                     "| s1   0    c1    c1·d3 |\n"
                     "|  0  -1     0   d1+d2  |\n"
                     "|  0   0     0      1   |"
        }
    ]

    # Bucle para insertar todas las secciones
    for slide in diapositivas:
        h = doc.add_heading(slide["titulo"], level=1)
        
        parrafo = doc.add_paragraph()
        for linea in slide["texto"].split('\n'):
            run = parrafo.add_run(linea + "\n")
            
            # Formateo visual
            if "[INSERTAR IMAGEN" in linea:
                run.font.color.rgb = RGBColor(220, 38, 38) # Rojo para las etiquetas
                run.font.bold = True
                run.font.name = 'Google Sans'
            
            # Condición estricta para formatear TODAS las matrices y tablas con Consolas
            elif "|" in linea or "Artic." in linea or "   1   |" in linea or "   2   |" in linea or "   3   |" in linea:
                run.font.name = 'Consolas'
                run.font.size = Pt(10) # Letra apenas más chica para que las matrices no salten de renglón
            else:
                run.font.name = 'Google Sans'
                run.font.size = Pt(11)

    # Guardar el MEGA archivo
    doc.save('Mega_Apunte_Robotica_Completo.docx')
    print("¡Éxito total! El mega-documento 'Mega_Apunte_Robotica_Completo.docx' ha sido generado con las 23 diapositivas íntegras.")

if __name__ == '__main__':
    crear_apunte_masivo()